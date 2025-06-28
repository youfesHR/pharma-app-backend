import os
import json
import logging
from datetime import datetime
import gspread
import requests
import docx
from io import BytesIO
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS

# --- SETUP ---
logging.basicConfig(level=logging.INFO)
app = Flask(__name__)
CORS(app) 

# --- CONFIGURATION & CREDENTIALS ---
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
# We load the credentials once and store them.
try:
    GCP_CREDENTIALS = json.loads(os.environ.get('GCP_CREDENTIALS_JSON'))
except Exception as e:
    app.logger.error(f"FATAL: Could not parse GCP_CREDENTIALS_JSON. Check the environment variable. Error: {e}")
    GCP_CREDENTIALS = None

# --- DATABASE HELPER ---
# This function connects to a specific sheet *when needed*. This is more robust.
def get_sheet(worksheet_name):
    if not GCP_CREDENTIALS:
        raise Exception("GCP credentials are not loaded.")
    
    try:
        gc = gspread.service_account_from_dict(GCP_CREDENTIALS)
        spreadsheet = gc.open("PharmaFeedbackApp")
        return spreadsheet.worksheet(worksheet_name)
    except gspread.exceptions.SpreadsheetNotFound:
        raise Exception(f"Spreadsheet 'PharmaFeedbackApp' not found or not shared.")
    except gspread.exceptions.WorksheetNotFound:
        raise Exception(f"Worksheet '{worksheet_name}' not found in the spreadsheet.")
    except Exception as e:
        raise Exception(f"A gspread error occurred: {e}")

# --- GEMINI HELPER FUNCTIONS (UNCHANGED) ---
def analyze_with_gemini(text):
    # ... (code is identical to your version)
    if not GEMINI_API_KEY:
        return {"category": "Config Error", "sentiment": 0, "error": "GEMINI_API_KEY not set."}
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key={GEMINI_API_KEY}"
    prompt = f"""Analyze the following customer feedback. Provide a JSON object with two keys: "category" and "sentiment". The "category" must be one of the following: [Packaging, Formula, Color, Smell, Efficacy, Side Effect, Price, Documentation, Other]. The "sentiment" must be a number between -1.0 (very negative) and 1.0 (very positive). Feedback to analyze: \"{text}\""""
    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    headers = {'Content-Type': 'application/json'}
    try:
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()
        json_response = response.json()
        model_response_text = json_response['candidates'][0]['content']['parts'][0]['text']
        clean_json_string = model_response_text.replace('```json', '').replace('```', '').strip()
        analysis = json.loads(clean_json_string)
        return {"category": analysis.get('category', 'Parse Error'), "sentiment": analysis.get('sentiment', 0), "error": ""}
    except Exception as e:
        app.logger.error(f"Gemini API Error: {e}", exc_info=True)
        return {"category": "AI_Error", "sentiment": 0, "error": str(e)}

def generate_text_with_gemini(prompt):
    # ... (code is identical to your version)
    if not GEMINI_API_KEY:
        return "Error: GEMINI_API_KEY not set."
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key={GEMINI_API_KEY}"
    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    headers = {'Content-Type': 'application/json'}
    try:
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()
        json_response = response.json()
        return json_response['candidates'][0]['content']['parts'][0]['text']
    except Exception as e:
        app.logger.error(f"Gemini text generation error: {e}", exc_info=True)
        return "Error: Could not generate report text from AI."


# --- API ENDPOINTS (ROUTES) ---

@app.route('/health', methods=['GET'])
def health_check():
    # This is a new endpoint for debugging.
    try:
        # Try to connect to a sheet to test the whole connection.
        get_sheet("Feedback") 
        return jsonify({"status": "ok", "message": "Server is running and can connect to Google Sheets."})
    except Exception as e:
        app.logger.error(f"Health check failed: {str(e)}")
        return jsonify({"status": "error", "message": f"Server is running, but database connection failed: {str(e)}"}), 500

@app.route('/get-products', methods=['GET'])
def get_products():
    try:
        products_sheet = get_sheet("Products")
        products = products_sheet.col_values(1)[1:] 
        return jsonify({"status": "success", "products": products})
    except Exception as e:
        app.logger.error("An error occurred in /get-products", exc_info=True)
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/submit-feedback', methods=['POST'])
def submit_feedback():
    try:
        feedback_sheet = get_sheet("Feedback")
        data = request.json
        ai_analysis = analyze_with_gemini(f"Feedback: {data['feedbackText']}. Suggestion: {data['suggestionText']}")
        timestamp = datetime.utcnow().isoformat() + "Z" 
        new_row = [timestamp, data.get('productName'), data.get('feedbackText'), data.get('suggestionText'), data.get('clientName'), data.get('clientEmail'), ai_analysis['category'], ai_analysis['sentiment'], ai_analysis['error']]
        feedback_sheet.append_row(new_row)
        return jsonify({"status": "success", "message": "Feedback submitted."})
    except Exception as e:
        app.logger.error("An error occurred in /submit-feedback", exc_info=True)
        return jsonify({"status": "error", "message": str(e)}), 500
        
@app.route('/admin-login', methods=['POST'])
def admin_login():
    try:
        admin_sheet = get_sheet("AdminUsers")
        data = request.json
        users = admin_sheet.get_all_records()
        for user in users:
            if user['Username'] == data['username'] and str(user['Password']) == data['password']:
                return jsonify({"status": "success", "message": "Login successful."})
        return jsonify({"status": "error", "message": "Invalid credentials."}), 401
    except Exception as e:
        app.logger.error("An error occurred in /admin-login", exc_info=True)
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/get-all-feedback', methods=['GET'])
def get_all_feedback():
    try:
        feedback_sheet = get_sheet("Feedback")
        all_feedback = feedback_sheet.get_all_records()
        return jsonify({"status": "success", "feedback": all_feedback})
    except Exception as e:
        app.logger.error("An error occurred in /get-all-feedback", exc_info=True)
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/generate-report', methods=['GET'])
def generate_report():
    try:
        feedback_sheet = get_sheet("Feedback")
        lang = request.args.get('lang', 'english')
        all_feedback = feedback_sheet.get_all_records()
        if not all_feedback:
            return jsonify({"status": "error", "message": "No feedback data to generate a report."}), 404
        # ... (rest of report generation is identical to your version)
        data_summary_for_prompt = "\n".join([str(item) for item in all_feedback])
        report_prompt = f"""You are a senior business analyst for a pharmaceutical company. Your task is to write a concise, professional executive summary report based on the following raw customer feedback data. IMPORTANT: The entire report must be written in {lang}. The report should include these sections: 1. **Overall Summary:** A brief, high-level overview of the findings. 2. **Key Positive Themes:** What are customers consistently happy about? 3. **Key Areas for Improvement:** What are the most common complaints? Group similar issues. 4. **Actionable Recommendations:** Suggest 2-3 specific, concrete actions the company should take. Do not just list the data. Synthesize it into an insightful report in {lang}. --- RAW DATA --- {data_summary_for_prompt} --- END OF RAW DATA --- """
        generated_report_text = generate_text_with_gemini(report_prompt)
        document = docx.Document()
        document.add_heading('Customer Feedback Report', level=0)
        for paragraph in generated_report_text.split('\n'):
            if paragraph.strip():
                document.add_paragraph(paragraph)
        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return send_file( file_stream, as_attachment=True, download_name=f'Pharma_Feedback_Report_{lang}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        app.logger.error("An error occurred in /generate-report", exc_info=True)
        return jsonify({"status": "error", "message": "Could not generate report."}), 500

# --- START THE SERVER ---
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port)